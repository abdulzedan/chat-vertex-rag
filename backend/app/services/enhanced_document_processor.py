import csv
import logging
import os
import re
import unicodedata
from collections import Counter
from dataclasses import dataclass
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple

import nltk
import openpyxl
import PyPDF2
from docx import Document as DocxDocument
from nltk.corpus import stopwords
from nltk.tokenize import sent_tokenize, word_tokenize
from vertexai.generative_models import GenerativeModel, Part

# Optional Document AI import
try:
    from app.services.document_ai_processor import DocumentAIProcessor

    DOCUMENT_AI_AVAILABLE = True
except ImportError:
    DOCUMENT_AI_AVAILABLE = False
    DocumentAIProcessor = None

# Download required NLTK data
try:
    nltk.data.find("tokenizers/punkt_tab")
except LookupError:
    try:
        nltk.download("punkt_tab")
    except Exception:
        # Fallback to older punkt if punkt_tab fails
        nltk.download("punkt")

try:
    nltk.data.find("corpora/stopwords")
except LookupError:
    nltk.download("stopwords")

logger = logging.getLogger(__name__)

# Constants for text processing
MIN_LINE_LENGTH = 3
MIN_SECTION_LENGTH = 5
MAX_SECTION_LENGTH = 50
MIN_FALSE_POSITIVE_LENGTH = 3
STRUCTURED_CONTENT_THRESHOLD = 0.2
GEMINI_FILE_SIZE_LIMIT_MB = 19
MIN_CHUNK_SENTENCES = 5
OVERLAP_SENTENCES = 2
KEYWORD_LIMIT = 8
TOKEN_ESTIMATE_DIVISOR = 4


@dataclass
class DocumentChunk:
    """Enhanced chunk with metadata"""

    text: str
    chunk_id: int
    start_char: int
    end_char: int
    metadata: Dict[str, Any]


@dataclass
class DocumentMetadata:
    """Document-level metadata"""

    filename: str
    file_type: str
    page_count: int
    word_count: int
    char_count: int
    has_tables: bool
    has_images: bool
    language: str
    creation_date: Optional[datetime]
    sections: List[str]
    extracted_entities: Dict[str, List[str]]


class EnhancedDocumentProcessor:
    """Advanced document processor with semantic chunking and metadata extraction"""

    def __init__(self):
        self.vision_model = GenerativeModel("gemini-2.0-flash-001")
        # Optimized chunking parameters for better context preservation
        self.min_chunk_size = 300  # Minimum viable context
        self.max_chunk_size = 1500  # Optimal for completeness without fragmentation
        self.chunk_overlap = 150  # Substantial overlap to maintain context
        try:
            self.stop_words = set(stopwords.words("english"))
        except LookupError:
            nltk.download("stopwords")
            self.stop_words = set(stopwords.words("english"))

        # Optional Document AI processor
        use_document_ai_env = os.getenv("USE_DOCUMENT_AI", "false")
        use_gemini_fallback_env = os.getenv("USE_GEMINI_FALLBACK", "true")

        logger.info(f"USE_DOCUMENT_AI environment variable: {use_document_ai_env}")
        logger.info(
            f"USE_GEMINI_FALLBACK environment variable: {use_gemini_fallback_env}"
        )
        logger.info(f"DOCUMENT_AI_AVAILABLE: {DOCUMENT_AI_AVAILABLE}")

        self.use_document_ai = (
            DOCUMENT_AI_AVAILABLE and use_document_ai_env.lower() == "true"
        )
        self.use_gemini_fallback = use_gemini_fallback_env.lower() == "true"

        if self.use_document_ai:
            self.document_ai = DocumentAIProcessor()
            logger.info("Document AI Layout Parser enabled")
        else:
            self.document_ai = None
            logger.info(
                f"Using standard document processing (Document AI available: {DOCUMENT_AI_AVAILABLE}, env var: {use_document_ai_env})"
            )

        if self.use_gemini_fallback:
            logger.info("Gemini multimodal fallback enabled")
        else:
            logger.info(
                "Gemini multimodal fallback disabled - will skip to standard processing"
            )

    def _preprocess_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove excessive whitespace
        text = re.sub(r"\s+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)

        # Fix common OCR errors
        text = self._fix_ocr_errors(text)

        # Normalize unicode characters
        text = unicodedata.normalize("NFKD", text)

        # Remove control characters
        text = "".join(
            char
            for char in text
            if unicodedata.category(char)[0] != "C" or char == "\n"
        )

        return text.strip()

    def _fix_ocr_errors(self, text: str) -> str:
        """Fix common OCR errors"""
        corrections = {
            r"\bl\s+(\w)": r"I \1",  # l -> I
            r"(\w)\s+l\b": r"\1 I",  # l -> I
            r"\b0\s+(\w)": r"O \1",  # 0 -> O
            r"(\w)\s+0\b": r"\1 O",  # 0 -> O
            r"ﬁ": "fi",
            r"ﬂ": "fl",
            r"–": "-",
            r'"': '"',
            r""": "'",
            r""": "'",
        }

        for pattern, replacement in corrections.items():
            text = re.sub(pattern, replacement, text)

        return text

    def _extract_metadata(
        self,
        text: str,
        file_type: str,
        filename: str,
        additional_info: Dict[str, Any] = None,
    ) -> DocumentMetadata:
        """Extract comprehensive metadata from document"""
        # Basic counts
        words = word_tokenize(text)
        word_count = len(words)
        char_count = len(text)

        # Detect language (simplified - could use langdetect)
        language = "en"  # Default to English

        # Extract sections/headers
        sections = self._extract_sections(text)

        # Extract entities
        entities = self._extract_entities(text)

        # Check for tables and images
        has_tables = bool(re.search(r"\|.*\|", text) or "table" in text.lower())
        has_images = bool(additional_info and additional_info.get("has_images", False))

        # Page count
        page_count = additional_info.get("page_count", 1) if additional_info else 1

        return DocumentMetadata(
            filename=filename,
            file_type=file_type,
            page_count=page_count,
            word_count=word_count,
            char_count=char_count,
            has_tables=has_tables,
            has_images=has_images,
            language=language,
            creation_date=None,
            sections=sections,
            extracted_entities=entities,
        )

    def _extract_sections(self, text: str) -> List[str]:
        """Extract section headers from text with improved patterns for PDFs"""
        sections = []

        # Enhanced header patterns for PDF content
        header_patterns = [
            r"^#{1,6}\s+(.+)$",  # Markdown headers
            r"^([A-Z][A-Z\s]{2,}):?\s*$",  # ALL CAPS headers (at least 3 chars)
            r"^\d+\.?\s+([A-Z].{3,})$",  # Numbered sections
            r"^[IVX]+\.\s+(.+)$",  # Roman numeral sections
            r"^([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*):?\s*$",  # Title Case headers
            r"^-\s*([A-Z].{3,})$",  # Dash bullet points
            r"^•\s*([A-Z].{3,})$",  # Bullet points
            r"^\*\s*([A-Z].{3,})$",  # Asterisk bullet points
            r"^([A-Z].{10,})\s*:$",  # Long phrases ending with colon
        ]

        lines = text.split("\n")
        for raw_line in lines:
            line = raw_line.strip()
            if len(line) < MIN_LINE_LENGTH:  # Skip very short lines
                continue

            for pattern in header_patterns:
                match = re.match(pattern, line)
                if match:
                    header_text = match.group(1) if match.lastindex else match.group(0)
                    # Filter out common false positives
                    if not self._is_false_positive_header(header_text):
                        sections.append(header_text.strip())
                    break

        # Also look for standalone capitalized words/phrases
        for raw_line in lines:
            line = raw_line.strip()
            if (
                len(line) > MIN_SECTION_LENGTH
                and len(line) < MAX_SECTION_LENGTH
                and line.isupper()
                and not any(char.isdigit() for char in line)
                and line not in sections
            ):
                sections.append(line)

        return sections[:20]  # Limit to top 20 sections

    def _is_false_positive_header(self, text: str) -> bool:
        """Check if text is likely a false positive header"""
        false_positives = [
            "PAGE",
            "TOTAL",
            "AMOUNT",
            "DATE",
            "NAME",
            "ADDRESS",
            "PHONE",
            "EMAIL",
            "SIGNATURE",
            "NOTES",
            "COMMENTS",
            "DESCRIPTION",
        ]
        return text.upper() in false_positives or len(text) < MIN_FALSE_POSITIVE_LENGTH

    def _extract_entities(self, text: str) -> Dict[str, List[str]]:
        """Extract named entities from text with enhanced patterns"""
        entities = {
            "dates": [],
            "emails": [],
            "urls": [],
            "phone_numbers": [],
            "percentages": [],
            "currency": [],
            "addresses": [],
            "abbreviations": [],
        }

        # Enhanced date patterns
        date_patterns = [
            r"\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b",
            r"\b\d{4}[/-]\d{1,2}[/-]\d{1,2}\b",
            r"\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b",
            r"\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\.?\s+\d{1,2},?\s+\d{4}\b",
        ]
        for pattern in date_patterns:
            entities["dates"].extend(re.findall(pattern, text, re.IGNORECASE))

        # Email pattern
        entities["emails"] = re.findall(
            r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b", text
        )

        # URL pattern
        entities["urls"] = re.findall(r"https?://[^\s]+", text)

        # Phone number patterns
        phone_patterns = [
            r"\b\d{3}[-.]?\d{3}[-.]?\d{4}\b",
            r"\(\d{3}\)\s*\d{3}[-.]?\d{4}\b",
            r"\b\d{3}\s+\d{3}\s+\d{4}\b",
        ]
        for pattern in phone_patterns:
            entities["phone_numbers"].extend(re.findall(pattern, text))

        # Percentage patterns
        entities["percentages"] = re.findall(r"\b\d+\.?\d*\s*%", text)

        # Currency patterns
        currency_patterns = [
            r"\$\d+(?:,\d{3})*(?:\.\d{2})?",
            r"\b\d+(?:,\d{3})*(?:\.\d{2})?\s*(?:dollars?|USD|CAD)\b",
        ]
        for pattern in currency_patterns:
            entities["currency"].extend(re.findall(pattern, text, re.IGNORECASE))

        # Address patterns (basic)
        address_patterns = [
            r"\b\d+\s+[A-Za-z\s]+(?:Street|St|Avenue|Ave|Road|Rd|Boulevard|Blvd|Drive|Dr|Lane|Ln)\b",
            r"\b[A-Za-z\s]+,\s*[A-Z]{2}\s+\d{5}(?:-\d{4})?\b",  # City, State ZIP
        ]
        for pattern in address_patterns:
            entities["addresses"].extend(re.findall(pattern, text, re.IGNORECASE))

        # Enhanced abbreviations (2-5 capital letters)
        entities["abbreviations"] = list(set(re.findall(r"\b[A-Z]{2,5}\b", text)))

        # Filter out common false positives
        entities["abbreviations"] = [
            abbr
            for abbr in entities["abbreviations"]
            if abbr
            not in [
                "THE",
                "AND",
                "FOR",
                "YOU",
                "ARE",
                "NOT",
                "BUT",
                "CAN",
                "ALL",
                "ANY",
                "NEW",
                "GET",
                "NOW",
                "MAY",
                "USE",
            ]
        ]

        # Limit entities and remove duplicates
        for key in entities:
            entities[key] = list(set(entities[key]))[:10]

        return entities

    def _detect_section_hint(
        self, chunk_text: str, sections: List[str]
    ) -> Optional[str]:
        """Try to align a chunk with the nearest known section heading."""
        if not sections:
            return None

        lowered_chunk = chunk_text.lower()
        for section in sections:
            if section and section.lower() in lowered_chunk:
                return section

        # Fallback to the leading line if no explicit match is found
        for line in chunk_text.split("\n"):
            clean_line = line.strip().strip("-•*")
            if not clean_line:
                continue
            if clean_line.lower().startswith("page "):
                continue
            return clean_line[:200]

        return None

    def _infer_page_range(self, chunk_text: str) -> Tuple[Optional[int], Optional[int]]:
        """Infer page range from embedded page markers."""
        matches = re.findall(r"--- Page (\d+) ---", chunk_text)
        if not matches:
            return None, None

        pages = [int(num) for num in matches]
        return min(pages), max(pages)

    def _extract_keywords(self, chunk_text: str) -> List[str]:
        """Return top keywords by frequency, excluding stopwords and short tokens."""
        try:
            tokens = word_tokenize(chunk_text)
        except LookupError:
            tokens = re.findall(r"\b\w+\b", chunk_text)

        filtered = [
            token.lower()
            for token in tokens
            if token.isalpha()
            and token.lower() not in self.stop_words
            and len(token) > 3
        ]

        if not filtered:
            return []

        counts = Counter(filtered)
        # Preserve deterministic order by frequency then alphabetically
        most_common = counts.most_common(KEYWORD_LIMIT)
        return [word for word, _ in most_common]

    def _derive_headline(self, chunk_text: str) -> Optional[str]:
        """Generate a concise headline from the first meaningful line."""
        for line in chunk_text.split("\n"):
            clean_line = line.strip()
            if not clean_line:
                continue
            if clean_line.startswith("--- Page"):
                continue
            words = clean_line.split()
            if not words:
                continue
            return " ".join(words[:12])
        return None

    def _build_chunk_metadata(
        self,
        *,
        chunk_text: str,
        doc_metadata: DocumentMetadata,
        source_start: int,
        source_end: int,
        content_type: str,
        overlap_with_previous: bool,
        contains_table: bool = False,
        extra: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        """Create a consistent metadata payload for chunk indexing."""

        section_hint = self._detect_section_hint(chunk_text, doc_metadata.sections)
        page_start, page_end = self._infer_page_range(chunk_text)
        keywords = self._extract_keywords(chunk_text)
        entities = self._extract_entities(chunk_text)

        metadata: Dict[str, Any] = {
            "filename": doc_metadata.filename,
            "content_type": content_type,
            "section_hint": section_hint,
            "headline": self._derive_headline(chunk_text),
            "page_start": page_start,
            "page_end": page_end,
            "keyword_terms": keywords,
            "entities": entities,
            "word_count": len(chunk_text.split()),
            "char_count": len(chunk_text),
            "token_estimate": max(1, len(chunk_text) // TOKEN_ESTIMATE_DIVISOR),
            "overlap_with_previous": overlap_with_previous,
            "contains_table": contains_table,
            "source_start": source_start,
            "source_end": source_end,
        }

        if extra:
            metadata.update(extra)

        return metadata

    def _simple_sentence_split(self, text: str) -> List[str]:
        """Simple sentence splitting fallback when NLTK is not available"""
        # Split on sentence endings followed by space and capital letter or end of string
        sentences = re.split(r"[.!?]+(?:\s+[A-Z]|\s*$)", text)
        # Clean up and filter empty sentences
        sentences = [s.strip() for s in sentences if s.strip()]
        return sentences

    def _has_structured_content(self, text: str) -> bool:
        """Detect if text contains structured content (tables, lists, etc.)"""
        # Check for table markers
        if "=== TABLE" in text or "|" in text:
            return True

        # Check for list patterns
        list_patterns = [
            r"^\s*\d+\.\s",  # Numbered lists
            r"^\s*[•\-\*]\s",  # Bullet points
            r"^\s*[A-Z]\.\s",  # Letter lists
        ]

        lines = text.split("\n")
        list_count = 0
        for line in lines:
            for pattern in list_patterns:
                if re.search(pattern, line, re.MULTILINE):
                    list_count += 1
                    break

        # If more than 20% of lines are list items, consider it structured
        if len(lines) > 0 and list_count / len(lines) > STRUCTURED_CONTENT_THRESHOLD:
            return True

        return False

    def _semantic_chunk(
        self, text: str, doc_metadata: DocumentMetadata
    ) -> List[DocumentChunk]:
        """Create semantic chunks based on sentence boundaries and structure with improved sizing"""
        chunks = []
        chunk_id = 0

        # Detect if content has structured data for adaptive chunking
        has_structured_content = self._has_structured_content(text)
        if has_structured_content:
            logger.info(
                "Detected structured content - using larger chunks to preserve context"
            )
            # Use larger chunks for structured content
            max_chunk_size = 2000
            min_chunk_size = 500
        else:
            max_chunk_size = self.max_chunk_size
            min_chunk_size = self.min_chunk_size

        # Split into sentences with fallback
        try:
            sentences = sent_tokenize(text)
        except LookupError:
            # Fallback to simple sentence splitting if NLTK data not available
            logger.warning(
                "NLTK punkt data not available, using simple sentence splitting"
            )
            sentences = self._simple_sentence_split(text)

        if not sentences:
            return []

        current_chunk = []
        current_length = 0
        start_char = 0

        current_chunk_has_overlap = False

        for i, sentence in enumerate(sentences):
            sentence_length = len(sentence)

            # Decide whether to start a new chunk before adding this sentence
            should_create_chunk = False
            split_reason = "continuation"

            if current_length + sentence_length > max_chunk_size and current_chunk:
                should_create_chunk = True
                split_reason = "max_size"

            elif current_length >= min_chunk_size and current_chunk:
                next_sentence_break = sentence.strip().endswith((".", "!", "?", ":"))
                if i < len(sentences) - 1 and next_sentence_break:
                    should_create_chunk = True
                    split_reason = "natural_break"
                elif len(current_chunk) >= MIN_CHUNK_SENTENCES:
                    should_create_chunk = True
                    split_reason = "sentence_limit"

            if should_create_chunk:
                chunk_text = " ".join(current_chunk)
                chunk_end = start_char + len(chunk_text)

                chunk_metadata = self._build_chunk_metadata(
                    chunk_text=chunk_text,
                    doc_metadata=doc_metadata,
                    source_start=start_char,
                    source_end=chunk_end,
                    content_type="text",
                    overlap_with_previous=current_chunk_has_overlap,
                    extra={
                        "sentence_count": len(current_chunk),
                        "split_reason": split_reason,
                        "chunk_index": chunk_id,
                    },
                )

                chunks.append(
                    DocumentChunk(
                        text=chunk_text,
                        chunk_id=chunk_id,
                        start_char=start_char,
                        end_char=chunk_end,
                        metadata=chunk_metadata,
                    )
                )
                chunk_id += 1

                # Prepare overlap for the next chunk
                overlap_count = min(OVERLAP_SENTENCES, len(current_chunk))
                overlap_sentences = current_chunk[-overlap_count:]
                overlap_text = " ".join(overlap_sentences)

                start_char = (
                    chunk_end - len(overlap_text) if overlap_text else chunk_end
                )

                current_chunk = overlap_sentences + [sentence]
                current_length = sum(len(s) for s in current_chunk)
                current_chunk_has_overlap = bool(overlap_sentences)
            else:
                current_chunk.append(sentence)
                current_length += sentence_length

        if current_chunk:
            chunk_text = " ".join(current_chunk)
            chunk_end = start_char + len(chunk_text)

            chunk_metadata = self._build_chunk_metadata(
                chunk_text=chunk_text,
                doc_metadata=doc_metadata,
                source_start=start_char,
                source_end=chunk_end,
                content_type="text",
                overlap_with_previous=current_chunk_has_overlap,
                extra={
                    "sentence_count": len(current_chunk),
                    "split_reason": "final",
                    "chunk_index": chunk_id,
                },
            )

            chunks.append(
                DocumentChunk(
                    text=chunk_text,
                    chunk_id=chunk_id,
                    start_char=start_char,
                    end_char=chunk_end,
                    metadata=chunk_metadata,
                )
            )

        return chunks

    def _table_aware_chunk(
        self, text: str, doc_metadata: DocumentMetadata, tables: List[Dict[str, Any]]
    ) -> List[DocumentChunk]:
        """Create chunks that preserve table integrity - tables are never split"""
        chunks = []
        chunk_id = 0

        # Identify table boundaries in the text
        table_boundaries = []
        for i, _table in enumerate(tables):
            # Find where each table appears in the text
            table_marker = f"=== TABLE {i + 1} ==="
            start_pos = text.find(table_marker)
            if start_pos != -1:
                # Find the end of this table (start of next table or end of text)
                next_table_marker = f"=== TABLE {i + 2} ==="
                end_pos = text.find(next_table_marker, start_pos)
                if end_pos == -1:
                    # This is the last table, find end by looking for non-table content
                    end_pos = text.find("\n\n", start_pos + len(table_marker) + 100)
                    if end_pos == -1:
                        end_pos = len(text)

                table_boundaries.append((start_pos, end_pos, i))

        # Sort table boundaries by position
        table_boundaries.sort(key=lambda x: x[0])

        # Split text into segments: [text_before_table1, table1, text_between_tables, table2, etc.]
        segments = []
        last_end = 0

        for start_pos, end_pos, table_idx in table_boundaries:
            # Add text before this table
            if start_pos > last_end:
                text_segment = text[last_end:start_pos].strip()
                if text_segment:
                    segments.append(("text", text_segment))

            # Add the table as one complete segment
            table_segment = text[start_pos:end_pos].strip()
            segments.append(("table", table_segment, table_idx))
            last_end = end_pos

        # Add any remaining text after the last table
        if last_end < len(text):
            remaining_text = text[last_end:].strip()
            if remaining_text:
                segments.append(("text", remaining_text))

        # Create chunks from segments
        start_char = 0
        for segment in segments:
            if segment[0] == "table":
                # Table segment - always create as single chunk
                table_text = segment[1]
                table_idx = segment[2]

                chunk_end = start_char + len(table_text)
                chunk_metadata = self._build_chunk_metadata(
                    chunk_text=table_text,
                    doc_metadata=doc_metadata,
                    source_start=start_char,
                    source_end=chunk_end,
                    content_type="table",
                    overlap_with_previous=False,
                    contains_table=True,
                    extra={
                        "table_index": table_idx,
                        "is_complete_table": True,
                        "chunk_index": chunk_id,
                    },
                )

                chunks.append(
                    DocumentChunk(
                        text=table_text,
                        chunk_id=chunk_id,
                        start_char=start_char,
                        end_char=chunk_end,
                        metadata=chunk_metadata,
                    )
                )
                chunk_id += 1
                start_char += len(table_text)

            else:
                # Text segment - chunk normally but smaller to compensate for large table chunks
                text_segment = segment[1]
                text_chunks = self._chunk_text_segment(
                    text_segment, chunk_id, start_char, doc_metadata
                )
                chunks.extend(text_chunks)
                chunk_id += len(text_chunks)
                start_char += len(text_segment)

        logger.info(
            f"Table-aware chunking created {len(chunks)} chunks ({len([c for c in chunks if c.metadata.get('contains_table')])} table chunks)"
        )
        return chunks

    def _chunk_text_segment(
        self,
        text: str,
        start_chunk_id: int,
        start_char: int,
        doc_metadata: DocumentMetadata,
    ) -> List[DocumentChunk]:
        """Chunk a text segment (non-table content) with optimized sizing"""
        chunks: List[DocumentChunk] = []

        max_chunk_size = 1200
        min_chunk_size = self.min_chunk_size

        try:
            sentences = sent_tokenize(text)
        except Exception:
            sentences = self._simple_sentence_split(text)

        if not sentences:
            return []

        current_chunk: List[str] = []
        current_length = 0
        chunk_start = start_char
        chunk_id = start_chunk_id
        current_chunk_has_overlap = False

        for i, sentence in enumerate(sentences):
            sentence_length = len(sentence)
            should_create_chunk = False
            split_reason = "continuation"

            if current_length + sentence_length > max_chunk_size and current_chunk:
                should_create_chunk = True
                split_reason = "max_size"
            elif current_length >= min_chunk_size and current_chunk:
                next_sentence_break = sentence.strip().endswith((".", "!", "?", ":"))
                if i < len(sentences) - 1 and next_sentence_break:
                    should_create_chunk = True
                    split_reason = "natural_break"
                elif len(current_chunk) >= MIN_CHUNK_SENTENCES:
                    should_create_chunk = True
                    split_reason = "sentence_limit"

            if should_create_chunk:
                chunk_text = " ".join(current_chunk)
                chunk_end = chunk_start + len(chunk_text)

                chunk_metadata = self._build_chunk_metadata(
                    chunk_text=chunk_text,
                    doc_metadata=doc_metadata,
                    source_start=chunk_start,
                    source_end=chunk_end,
                    content_type="text",
                    overlap_with_previous=current_chunk_has_overlap,
                    extra={
                        "sentence_count": len(current_chunk),
                        "split_reason": split_reason,
                        "chunk_index": chunk_id,
                    },
                )

                chunks.append(
                    DocumentChunk(
                        text=chunk_text,
                        chunk_id=chunk_id,
                        start_char=chunk_start,
                        end_char=chunk_end,
                        metadata=chunk_metadata,
                    )
                )
                chunk_id += 1

                overlap_count = min(OVERLAP_SENTENCES, len(current_chunk))
                overlap_sentences = current_chunk[-overlap_count:]
                overlap_text = " ".join(overlap_sentences)

                chunk_start = (
                    chunk_end - len(overlap_text) if overlap_text else chunk_end
                )
                current_chunk = overlap_sentences + [sentence]
                current_length = sum(len(s) for s in current_chunk)
                current_chunk_has_overlap = bool(overlap_sentences)
            else:
                current_chunk.append(sentence)
                current_length += sentence_length

        if current_chunk:
            chunk_text = " ".join(current_chunk)
            chunk_end = chunk_start + len(chunk_text)

            chunk_metadata = self._build_chunk_metadata(
                chunk_text=chunk_text,
                doc_metadata=doc_metadata,
                source_start=chunk_start,
                source_end=chunk_end,
                content_type="text",
                overlap_with_previous=current_chunk_has_overlap,
                extra={
                    "sentence_count": len(current_chunk),
                    "split_reason": "final",
                    "chunk_index": chunk_id,
                },
            )

            chunks.append(
                DocumentChunk(
                    text=chunk_text,
                    chunk_id=chunk_id,
                    start_char=chunk_start,
                    end_char=chunk_end,
                    metadata=chunk_metadata,
                )
            )

        return chunks

    def _handle_tables(self, text: str) -> Tuple[str, List[Dict[str, Any]]]:
        """Extract and process tables from text"""
        tables = []
        table_pattern = r"(\|[^\n]+\|(?:\n\|[^\n]+\|)*)"

        def replace_table(match):
            table_text = match.group(0)
            table_id = len(tables)

            # Parse table
            rows = table_text.strip().split("\n")
            parsed_rows = []

            for row in rows:
                cells = [cell.strip() for cell in row.split("|")[1:-1]]
                parsed_rows.append(cells)

            tables.append(
                {"id": table_id, "rows": parsed_rows, "original_text": table_text}
            )

            # Return placeholder
            return f"\n[TABLE_{table_id}]\n"

        # Replace tables with placeholders
        processed_text = re.sub(table_pattern, replace_table, text)

        return processed_text, tables

    async def _process_pdf_with_gemini(
        self, file_path: str, filename: str
    ) -> Tuple[str, List[Dict], Dict]:
        """Process PDF using Gemini's multimodal capabilities for faster extraction"""
        try:
            logger.info(f"Processing PDF with Gemini multimodal: {filename}")

            # Read PDF file
            with open(file_path, "rb") as f:
                pdf_bytes = f.read()

            # Check file size - Gemini has a 20MB limit for inline data
            file_size_mb = len(pdf_bytes) / (1024 * 1024)
            if file_size_mb > GEMINI_FILE_SIZE_LIMIT_MB:  # Leave some buffer
                logger.warning(
                    f"PDF too large for Gemini ({file_size_mb:.1f}MB), falling back to standard processing"
                )
                return None, [], {}

            # Create Part object for Gemini
            pdf_part = Part.from_data(mime_type="application/pdf", data=pdf_bytes)

            # Optimized prompt for fast extraction
            prompt = """Extract all text content from this PDF document.

            Requirements:
            1. Extract ALL text content including headers, body text, and footnotes
            2. Preserve the document structure and formatting
            3. For tables: format them as markdown tables
            4. Include page numbers as "--- Page X ---" markers
            5. Extract any important metadata like dates, amounts, percentages

            Output the complete extracted text."""

            # Use Gemini 2.0 Flash for optimal speed
            response = self.vision_model.generate_content([prompt, pdf_part])

            if response.text:
                extracted_text = response.text
                logger.info(
                    f"Gemini extracted {len(extracted_text)} characters from PDF"
                )

                # Quick table detection
                tables = []
                if "|" in extracted_text and "---" in extracted_text:
                    # Extract markdown tables
                    table_pattern = r"\|[^\n]+\|(?:\n\|[-\s|]+\|)?(?:\n\|[^\n]+\|)+"
                    table_matches = re.findall(table_pattern, extracted_text)
                    for i, table_text in enumerate(table_matches):
                        tables.append(
                            {
                                "table_index": i,
                                "content": table_text,
                                "type": "markdown",
                            }
                        )

                # Extract metadata
                metadata = {
                    "extraction_method": "gemini_multimodal",
                    "has_tables": len(tables) > 0,
                    "table_count": len(tables),
                }

                return extracted_text, tables, metadata
            else:
                logger.warning("Gemini returned no text")
                return None, [], {}

        except Exception as e:
            logger.warning(f"Gemini PDF processing failed: {e}")
            return None, [], {}

    async def _process_image_with_gemini(self, file_path: str, filename: str) -> str:
        """Process image using Gemini's vision capabilities for OCR"""
        try:
            logger.info(f"Processing image with Gemini vision: {filename}")

            # Read image file
            with open(file_path, "rb") as f:
                image_bytes = f.read()

            # Check file size - Gemini has a 20MB limit
            file_size_mb = len(image_bytes) / (1024 * 1024)
            if file_size_mb > GEMINI_FILE_SIZE_LIMIT_MB:
                logger.warning(f"Image too large for Gemini ({file_size_mb:.1f}MB)")
                return None

            # Determine MIME type
            mime_type = (
                "image/png" if filename.lower().endswith(".png") else "image/jpeg"
            )

            # Create Part object for Gemini
            image_part = Part.from_data(mime_type=mime_type, data=image_bytes)

            # Optimized prompt for OCR
            prompt = """Extract all text from this image with high accuracy.

            Requirements:
            1. Extract ALL visible text including headers, body text, labels, captions
            2. Preserve the reading order and structure as much as possible
            3. For tables or structured data: format as markdown tables
            4. Include any numbers, dates, or special characters exactly as shown
            5. If there's no text in the image, respond with "No text detected"

            Output only the extracted text."""

            response = self.vision_model.generate_content([prompt, image_part])

            if response.text and response.text.strip() != "No text detected":
                logger.info(
                    f"Gemini extracted {len(response.text)} characters from image"
                )
                return response.text
            else:
                logger.info("No text found in image")
                return ""

        except Exception as e:
            logger.warning(f"Gemini image processing failed: {e}")
            return None

    async def process_file(
        self, file_path: str, file_type: str, filename: str
    ) -> Dict[str, Any]:
        """Process a file with enhanced extraction and chunking"""
        try:
            logger.info(f"Processing file: {filename} (type: {file_type})")

            # Extract text based on file type
            additional_info = {}

            if file_type == "application/pdf":
                # Processing hierarchy:
                # 1. Document AI (if enabled) - best for complex documents
                # 2. Gemini multimodal - fast fallback
                # 3. PyPDF2 - basic fallback

                text = ""
                processed_successfully = False

                # Try Document AI first if enabled
                if self.use_document_ai and self.document_ai:
                    try:
                        logger.info("Processing PDF with Document AI Layout Parser")
                        doc_ai_result = await self.document_ai.process_document_online(
                            file_path, file_type, filename
                        )

                        # Convert Document AI result to our format
                        text = doc_ai_result["full_text"]

                        # Include table content in the text for better search
                        if doc_ai_result["tables"]:
                            logger.info(
                                f"Document AI found {len(doc_ai_result['tables'])} tables"
                            )
                            for i, table in enumerate(doc_ai_result["tables"]):
                                logger.info(
                                    f"Table {i}: {len(table.get('headers', []))} header rows, {len(table.get('rows', []))} data rows"
                                )

                                table_text = f"\n\n=== TABLE {i + 1} ===\n"

                                # Create proper markdown table
                                all_rows = []

                                # Add headers
                                if table.get("headers"):
                                    for header_row in table["headers"]:
                                        all_rows.append(header_row)
                                        logger.info(f"Header row: {header_row}")

                                # Add data rows
                                if table.get("rows"):
                                    for j, row in enumerate(table["rows"]):
                                        all_rows.append(row)
                                        if j < 5:  # Log first 5 rows
                                            logger.info(f"Data row {j}: {row}")

                                # Format as markdown table
                                if all_rows:
                                    # Determine max columns
                                    max_cols = (
                                        max(len(row) for row in all_rows)
                                        if all_rows
                                        else 0
                                    )

                                    # Pad rows to same length
                                    for row in all_rows:
                                        while len(row) < max_cols:
                                            row.append("")

                                    # First row as header
                                    if all_rows:
                                        table_text += (
                                            "| " + " | ".join(all_rows[0]) + " |\n"
                                        )
                                        table_text += (
                                            "|"
                                            + "|".join(
                                                [" --- " for _ in range(max_cols)]
                                            )
                                            + "|\n"
                                        )

                                        # Data rows
                                        for row in all_rows[1:]:
                                            table_text += (
                                                "| " + " | ".join(row) + " |\n"
                                            )

                                # Also add raw text version for search
                                table_text += "\nRaw table data:\n"
                                for row in all_rows:
                                    table_text += " ".join(row) + "\n"

                                text += table_text

                        additional_info = {
                            "page_count": doc_ai_result["metadata"]["page_count"],
                            "has_tables": doc_ai_result["metadata"]["has_tables"],
                            "document_ai_chunks": doc_ai_result["chunks"],
                            "tables": doc_ai_result["tables"],
                            "entities": doc_ai_result["metadata"]["entities"],
                            "document_ai_used": True,
                            "extraction_method": "document_ai",
                        }
                        processed_successfully = True
                        logger.info(
                            f"Document AI processing successful: {len(doc_ai_result['chunks'])} chunks, {len(doc_ai_result['tables'])} tables"
                        )
                    except Exception as e:
                        logger.warning(f"Document AI processing failed: {e}")

                # Fallback to Gemini multimodal if Document AI failed or not enabled AND Gemini fallback is enabled
                if not processed_successfully and self.use_gemini_fallback:
                    logger.info("Trying Gemini multimodal as fallback")
                    (
                        gemini_text,
                        gemini_tables,
                        gemini_metadata,
                    ) = await self._process_pdf_with_gemini(file_path, filename)

                    if gemini_text:
                        text = gemini_text
                        additional_info = {
                            "tables": gemini_tables,
                            "has_tables": len(gemini_tables) > 0,
                            "page_count": text.count("--- Page "),
                            "extraction_method": "gemini_multimodal",
                        }
                        additional_info.update(gemini_metadata)
                        processed_successfully = True
                        logger.info("Gemini multimodal processing successful")

                # Final fallback to PyPDF2
                if not processed_successfully:
                    logger.info("Falling back to PyPDF2 standard processing")
                    text, additional_info = await self._extract_pdf_text_enhanced(
                        file_path
                    )
                    additional_info["extraction_method"] = "pypdf2"
            elif file_type in ["image/png", "image/jpeg", "image/jpg"]:
                # Try Gemini vision first for better OCR if enabled
                if self.use_gemini_fallback:
                    text = await self._process_image_with_gemini(file_path, filename)
                    if text is not None:
                        additional_info["extraction_method"] = "gemini_vision"
                    else:
                        # Fallback to existing OCR method
                        logger.info(
                            "Gemini vision failed, falling back to standard image OCR"
                        )
                        text = await self._extract_image_text(file_path)
                        additional_info["extraction_method"] = "standard_ocr"
                else:
                    # Skip Gemini and go directly to standard OCR
                    logger.info("Using standard image OCR (Gemini fallback disabled)")
                    text = await self._extract_image_text(file_path)
                    additional_info["extraction_method"] = "standard_ocr"
            elif file_type == "text/csv":
                text = await self._extract_csv_text_enhanced(file_path)
            elif (
                file_type
                == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ):
                text, additional_info = await self._extract_docx_text_enhanced(
                    file_path
                )
            elif (
                file_type
                == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            ):
                text = await self._extract_excel_text(file_path)
            else:
                text = await self._extract_plain_text(file_path)

            if not text or len(text.strip()) < 10:
                raise ValueError(f"Could not extract meaningful text from {filename}")

            # Preprocess text
            text = self._preprocess_text(text)

            # Handle tables only if Document AI wasn't used
            if not additional_info.get("document_ai_used"):
                text, tables = self._handle_tables(text)
                additional_info["tables"] = tables
            else:
                # Document AI already extracted tables
                tables = additional_info.get("tables", [])
                logger.info(f"Using {len(tables)} tables from Document AI")

            # Extract metadata
            doc_metadata = self._extract_metadata(
                text, file_type, filename, additional_info
            )

            # Create semantic chunks - use table-aware chunking if Document AI was used
            if additional_info.get("document_ai_used"):
                chunks = self._table_aware_chunk(
                    text, doc_metadata, additional_info.get("tables", [])
                )
            else:
                chunks = self._semantic_chunk(text, doc_metadata)

            logger.info(
                f"Extracted {len(text)} characters, created {len(chunks)} semantic chunks"
            )

            chunk_records = [
                {
                    "id": chunk.chunk_id,
                    "text": chunk.text,
                    "start": chunk.start_char,
                    "end": chunk.end_char,
                    "metadata": chunk.metadata,
                }
                for chunk in chunks
            ]

            return {
                "filename": filename,
                "file_type": file_type,
                "full_text": text,
                "chunks": chunk_records,
                "chunk_details": chunk_records,
                "metadata": {
                    "page_count": doc_metadata.page_count,
                    "word_count": doc_metadata.word_count,
                    "char_count": doc_metadata.char_count,
                    "has_tables": doc_metadata.has_tables,
                    "has_images": doc_metadata.has_images,
                    "sections": doc_metadata.sections,
                    "entities": doc_metadata.extracted_entities,
                },
                "tables": tables,
                "chunk_count": len(chunks),
                "character_count": len(text),
            }

        except Exception as e:
            logger.error(f"Error processing file {filename}: {e}")
            raise

    async def _extract_pdf_text_enhanced(
        self, file_path: str
    ) -> Tuple[str, Dict[str, Any]]:
        """Enhanced PDF text extraction with structure preservation"""
        text = ""
        page_count = 0
        has_images = False

        try:
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                page_count = len(pdf_reader.pages)

                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            # Add page marker for better structure
                            text += f"\n\n--- Page {page_num + 1} ---\n\n"
                            text += page_text

                        # Check for images safely
                        try:
                            if hasattr(page, "images") and page.images:
                                has_images = True
                        except Exception:
                            # Ignore image detection errors
                            pass

                    except Exception as e:
                        logger.warning(
                            f"Could not extract text from page {page_num + 1}: {e}"
                        )
                        continue

        except Exception as e:
            logger.error(f"Error reading PDF file: {e}")
            raise

        if not text.strip():
            raise ValueError("No text could be extracted from the PDF")

        return text, {"page_count": page_count, "has_images": has_images}

    async def _extract_csv_text_enhanced(self, file_path: str) -> str:
        """Enhanced CSV extraction with better formatting"""
        try:
            text = "CSV Document\n\n"

            with open(file_path, encoding="utf-8") as file:
                # Detect delimiter
                sample = file.read(1024)
                file.seek(0)

                sniffer = csv.Sniffer()
                delimiter = sniffer.sniff(sample).delimiter

                reader = csv.reader(file, delimiter=delimiter)

                # Read headers
                headers = next(reader, None)
                if headers:
                    text += "Column Headers:\n"
                    text += " | ".join(headers) + "\n\n"

                    # Create table representation
                    text += "Data:\n"
                    text += "|" + "|".join(headers) + "|\n"
                    text += "|" + "|".join(["-" * len(h) for h in headers]) + "|\n"

                    # Read data rows
                    row_count = 0
                    for row in reader:
                        if row_count < 100:
                            if len(row) == len(headers):
                                text += "|" + "|".join(row) + "|\n"
                            row_count += 1
                        else:
                            remaining = sum(1 for _ in reader) + 1
                            text += f"\n... and {remaining} more rows\n"
                            break

            return text

        except Exception as e:
            logger.error(f"Error extracting CSV text: {e}")
            raise

    async def _extract_docx_text_enhanced(
        self, file_path: str
    ) -> Tuple[str, Dict[str, Any]]:
        """Enhanced Word document extraction with structure"""
        try:
            doc = DocxDocument(file_path)
            text = ""
            has_tables = len(doc.tables) > 0

            # Extract paragraphs with style information
            for paragraph in doc.paragraphs:
                if paragraph.style.name.startswith("Heading"):
                    # Preserve heading structure
                    level = (
                        int(paragraph.style.name[-1])
                        if paragraph.style.name[-1].isdigit()
                        else 1
                    )
                    text += "\n" + "#" * level + " " + paragraph.text + "\n\n"
                else:
                    text += paragraph.text + "\n"

            # Extract tables
            if has_tables:
                text += "\n\nTables:\n"
                for i, table in enumerate(doc.tables):
                    text += f"\nTable {i + 1}:\n"

                    # Get table data
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells]
                        table_data.append(row_data)

                    # Format as markdown table
                    if table_data:
                        # Headers
                        text += "|" + "|".join(table_data[0]) + "|\n"
                        text += (
                            "|"
                            + "|".join(["-" * len(cell) for cell in table_data[0]])
                            + "|\n"
                        )

                        # Data rows
                        for row in table_data[1:]:
                            text += "|" + "|".join(row) + "|\n"

                    text += "\n"

            return text, {"has_tables": has_tables}

        except Exception as e:
            logger.error(f"Error extracting Word document text: {e}")
            raise

    # Keep original methods for compatibility
    async def _extract_image_text(self, file_path: str) -> str:
        """Extract text from images using Gemini Vision"""
        try:
            with open(file_path, "rb") as f:
                image_data = f.read()

            image_part = Part.from_data(data=image_data, mime_type="image/jpeg")

            prompt = """Extract all text content from this image. Include:
            - All readable text, numbers, and data
            - Table contents formatted as markdown tables
            - Chart or graph labels and values
            - Any other textual information

            Preserve the structure and formatting where possible."""

            response = self.vision_model.generate_content([prompt, image_part])

            if response.text:
                return response.text
            else:
                raise ValueError("No text extracted from image")

        except Exception as e:
            logger.error(f"Error extracting text from image: {e}")
            raise

    async def _extract_excel_text(self, file_path: str) -> str:
        """Extract text from Excel files"""
        try:
            workbook = openpyxl.load_workbook(file_path)
            text = "Excel Document\n\n"

            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text += f"Sheet: {sheet_name}\n"

                # Get data from sheet
                data = []
                for row in sheet.iter_rows(values_only=True):
                    row_data = [str(cell) if cell is not None else "" for cell in row]
                    if any(cell.strip() for cell in row_data):
                        data.append(row_data)

                if data:
                    # Format as table
                    headers = data[0] if data else []
                    if headers:
                        text += "|" + "|".join(headers) + "|\n"
                        text += "|" + "|".join(["-" * len(h) for h in headers]) + "|\n"

                        for row in data[1:51]:  # Limit to 50 rows
                            text += "|" + "|".join(row) + "|\n"

                        if len(data) > 51:
                            text += f"\n... and {len(data) - 51} more rows\n"

                text += "\n"

            return text

        except Exception as e:
            logger.error(f"Error extracting Excel text: {e}")
            raise

    async def _extract_plain_text(self, file_path: str) -> str:
        """Extract text from plain text files"""
        try:
            encodings = ["utf-8", "latin-1", "cp1252"]

            for encoding in encodings:
                try:
                    with open(file_path, encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue

            raise ValueError("Could not decode file with any common encoding")

        except Exception as e:
            logger.error(f"Error extracting plain text: {e}")
            raise
